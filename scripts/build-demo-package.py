from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DEMO = ROOT / "demo-data"
DOCS = ROOT / "docs"
TMP = ROOT / "tmp" / "pdfs" / "demo-package"
OUTPUT_PDF = ROOT / "output" / "pdf"
FONT_CN = Path(r"C:\Windows\Fonts\simhei.ttf")
FONT_SERIF = Path(r"C:\Windows\Fonts\simsun.ttc")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def clean_inline(text: str) -> str:
    return re.sub(r"`([^`]*)`", r"\1", text).replace("**", "")


def new_decimal_num_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def set_run_font(run, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 ascii_font: str = "Calibri", east_asia: str = "微软雅黑") -> None:
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("基于填图对象智能识别的文本空间化算法模块")
    set_run_font(run, 8.5, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    prefix = footer.add_run("第 ")
    set_run_font(prefix, 8.5, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    suffix = footer.add_run(" 页")
    set_run_font(suffix, 8.5, MUTED)


def add_cover(doc: Document, title: str, subtitle: str, meta: str, pattern: str) -> None:
    if pattern == "editorial_cover":
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(82)
        kicker = doc.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(kicker.add_run("SYSTEM DEMONSTRATION GUIDE"), 10, BLUE, True)
        kicker.paragraph_format.space_after = Pt(18)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        set_run_font(p.add_run(title), 28, INK, True)
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(54)
        set_run_font(s.add_run(subtitle), 15, DARK_BLUE)
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        m.paragraph_format.space_after = Pt(4)
        set_run_font(m.add_run(meta), 10.5, MUTED)
    else:
        kicker = doc.add_paragraph()
        kicker.paragraph_format.space_after = Pt(8)
        set_run_font(kicker.add_run("LIVE DEMONSTRATION SCRIPT"), 10, BLUE, True)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        set_run_font(p.add_run(title), 27, INK, True)
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(20)
        set_run_font(s.add_run(subtitle), 14, DARK_BLUE)
        strip = doc.add_table(rows=1, cols=3)
        labels = [("15-18分钟", "建议时长"), ("9个环节", "完整流程"), ("证据可回查", "讲解重点")]
        for idx, (value, label) in enumerate(labels):
            cell = strip.cell(0, idx)
            set_cell_shading(cell, LIGHT_BLUE)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(cell.paragraphs[0].add_run(value + "\n"), 12, INK, True)
            set_run_font(cell.paragraphs[0].add_run(label), 8.5, MUTED)
        set_table_geometry(strip, [3120, 3120, 3120], indent_dxa=120)
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        m.paragraph_format.space_before = Pt(18)
        set_run_font(m.add_run(meta), 10, MUTED)
    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.rows[0].cells[0].paragraphs[0].paragraph_format.keep_with_next = True
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_run_font(p.add_run(clean_inline(value)), 9.5, INK if row_idx == 0 else None, row_idx == 0)
    set_repeat_table_header(table.rows[0])
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2100, 1500, 5760]
    elif cols == 4:
        widths = [1200, 3000, 1560, 3600]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), CALLOUT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(text), 10.5, INK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def append_inline_paragraph(doc: Document, text: str) -> None:
    text = clean_inline(text)
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    label_match = re.match(r"^(操作|讲稿|转场|验收点|说明|版本|建议时长|适用场景)：(.*)$", text)
    if label_match:
        set_run_font(p.add_run(label_match.group(1) + "："), 11, DARK_BLUE, True)
        set_run_font(p.add_run(label_match.group(2)), 11)
    else:
        set_run_font(p.add_run(text), 11)


def markdown_to_docx(markdown_path: Path, output_path: Path, title: str, subtitle: str,
                     meta: str, pattern: str) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc, title, subtitle, meta, pattern)
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    idx = 0
    skipped_title_headings = 0
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []
    active_number_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            append_inline_paragraph(doc, " ".join(line.strip() for line in paragraph_lines))
            paragraph_lines = []

    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                set_cell_shading(table.cell(0, 0), LIGHT_GRAY)
                p = table.cell(0, 0).paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, 9, INK, ascii_font="Consolas", east_asia="微软雅黑")
                set_table_geometry(table, [9360])
                code_lines = []
            in_code = not in_code
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not stripped:
            flush_paragraph()
            active_number_id = None
            idx += 1
            continue
        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[idx + 1]):
            flush_paragraph()
            active_number_id = None
            table_lines = [stripped]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            active_number_id = None
            level = len(heading.group(1))
            if level <= 2 and skipped_title_headings < 2:
                skipped_title_headings += 1
            else:
                doc.add_paragraph(heading.group(2), style=f"Heading {min(level, 3)}")
            idx += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            active_number_id = None
            add_callout(doc, stripped[1:].strip())
            idx += 1
            continue
        if re.match(r"^-\s+", stripped):
            flush_paragraph()
            active_number_id = None
            p = doc.add_paragraph(style="List Bullet")
            set_run_font(p.add_run(clean_inline(re.sub(r"^-\s+", "", stripped))), 11)
            idx += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            if active_number_id is None:
                active_number_id = new_decimal_num_id(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_number_id)
            set_run_font(p.add_run(clean_inline(numbered.group(1))), 11)
            idx += 1
            continue
        active_number_id = None
        paragraph_lines.append(stripped.replace("  ", " "))
        idx += 1
    flush_paragraph()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "GeoText Development Team"
    doc.save(output_path)


def build_drill_docx() -> None:
    doc = Document()
    configure_document(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("ZK001钻孔地质编录表（演示数据）"), 20, INK, True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run("铜绿山矿段 | WGS84 | 编制年份 2026"), 10.5, MUTED)

    add_table(doc, [
        ["字段", "记录值"],
        ["孔口坐标", "东经114.9384°、北纬30.0840°"],
        ["终孔深度", "286.50 m"],
        ["施工方向", "垂直孔"],
        ["调查区域", "湖北省大冶市铜绿山矿段"],
    ])
    doc.add_paragraph("钻孔分层记录", style="Heading 1")
    add_table(doc, [
        ["起止深度", "厚度", "地质描述"],
        ["0-38.20 m", "38.20 m", "第四系覆盖层，黄褐色黏土夹碎石。"],
        ["38.20-126.40 m", "88.20 m", "下三叠统大冶组中厚层状灰岩，局部发育矽卡岩化。"],
        ["126.40-188.70 m", "62.30 m", "燕山期闪长玢岩，侵入大冶组灰岩。"],
        ["188.70-214.30 m", "25.60 m", "Ⅰ号铜铁矿体，黄铜矿和磁铁矿呈浸染状、团块状分布。"],
        ["214.30-286.50 m", "72.20 m", "闪长玢岩与矽卡岩化灰岩互层。"],
    ])
    doc.add_paragraph("矿体与构造", style="Heading 1")
    add_callout(doc, "Ⅰ号铜铁矿体厚度25.60 m，铜平均品位1.26%，铁平均品位38.40%，赋存于矽卡岩化带，并受F1断裂控制。")
    append_inline_paragraph(doc, "F1断裂走向北东，倾向南东，倾角68°；在钻孔北西侧控制矿体展布。")
    append_inline_paragraph(doc, "说明：本编录表为系统功能演示而构造，不得用于生产决策。")
    path = DEMO / "02-ZK001钻孔地质编录表.docx"
    doc.core_properties.title = "ZK001钻孔地质编录表（演示数据）"
    doc.core_properties.author = "GeoText Development Team"
    doc.save(path)


def register_pdf_font() -> str:
    name = "SimHei"
    if name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(name, str(FONT_CN)))
    return name


def build_searchable_pdf() -> None:
    font = register_pdf_font()
    output = DEMO / "03-铜绿山矿段矿化特征简报.pdf"
    pdf = canvas.Canvas(str(output), pagesize=A4)
    width, height = A4
    pdf.setTitle("铜绿山矿段矿化特征简报（演示数据）")
    pdf.setAuthor("GeoText Development Team")
    pdf.setFont(font, 19)
    pdf.setFillColorRGB(0.04, 0.15, 0.27)
    pdf.drawString(56, height - 70, "铜绿山矿段矿化特征简报")
    pdf.setFont(font, 9)
    pdf.setFillColorRGB(0.4, 0.44, 0.5)
    pdf.drawString(56, height - 90, "可检索PDF演示资料 | 编制年份 2026 | 坐标系 WGS84")
    pdf.setStrokeColorRGB(0.18, 0.45, 0.7)
    pdf.line(56, height - 102, width - 56, height - 102)
    sections = [
        ("1  地质背景", [
            "铜绿山矿段位于湖北省大冶市，主要出露下三叠统大冶组中厚层状灰岩。",
            "燕山期闪长玢岩侵入大冶组灰岩，接触带发育矽卡岩化。",
        ]),
        ("2  矿化特征", [
            "一号铜铁矿体赋存于矽卡岩化带，走向延伸约420 m，厚度25.60 m。",
            "矿体铜平均品位1.26%，铁平均品位38.40%，主要矿物为黄铜矿和磁铁矿。",
        ]),
        ("3  控矿构造", [
            "F1断裂走向北东、倾向南东、倾角68°，控制一号铜铁矿体及矽卡岩型矿化。",
            "断裂由114.9310°E、30.0790°N延伸至114.9470°E、30.0890°N。",
        ]),
        ("4  结论", [
            "闪长玢岩侵入、矽卡岩化带和F1断裂共同构成矿体识别与空间化的重要证据链。",
        ]),
    ]
    y = height - 145
    for heading, paragraphs in sections:
        pdf.setFont(font, 13)
        pdf.setFillColorRGB(0.18, 0.45, 0.7)
        pdf.drawString(56, y, heading)
        y -= 27
        pdf.setFont(font, 10.5)
        pdf.setFillColorRGB(0.12, 0.15, 0.18)
        for paragraph in paragraphs:
            pdf.drawString(67, y, paragraph)
            y -= 24
        y -= 9
    pdf.setFont(font, 8.5)
    pdf.setFillColorRGB(0.45, 0.45, 0.45)
    pdf.drawString(56, 42, "说明：本资料为系统功能演示而构造，不代表真实勘查成果。")
    pdf.drawRightString(width - 56, 42, "第1页")
    pdf.save()
    OUTPUT_PDF.mkdir(parents=True, exist_ok=True)
    (OUTPUT_PDF / output.name).write_bytes(output.read_bytes())


def build_scanned_pdf() -> None:
    TMP.mkdir(parents=True, exist_ok=True)
    width_px, height_px = 2480, 3508
    image = Image.new("RGB", (width_px, height_px), "#f6f2e8")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(str(FONT_CN), 84)
    heading_font = ImageFont.truetype(str(FONT_CN), 52)
    body_font = ImageFont.truetype(str(FONT_CN), 45)
    small_font = ImageFont.truetype(str(FONT_CN), 34)
    draw.rectangle((120, 110, width_px - 120, height_px - 120), outline="#6e665a", width=4)
    draw.text((210, 190), "F1断裂野外地质记录卡", font=title_font, fill="#25221e")
    draw.line((200, 315, width_px - 200, 315), fill="#6e665a", width=4)
    fields = [
        "记录编号：TL-F1-2026-04",
        "调查位置：湖北省大冶市铜绿山矿段",
        "起点坐标：114.9310°E，30.0790°N",
        "终点坐标：114.9470°E，30.0890°N",
        "构造名称：F1断裂",
        "走向：北东    倾向：南东    倾角：68°",
    ]
    y = 410
    for line in fields:
        draw.text((220, y), line, font=body_font, fill="#25221e")
        y += 125
    draw.text((220, y + 30), "野外描述", font=heading_font, fill="#25221e")
    y += 125
    descriptions = [
        "断裂带宽约8-12 m，见角砾岩和碎裂岩。",
        "F1断裂控制一号铜铁矿体及矽卡岩型矿化，",
        "断裂两侧可见黄铜矿、磁铁矿浸染。",
        "建议沿北东方向追索，并核对ZK001钻孔资料。",
    ]
    for line in descriptions:
        draw.text((250, y), line, font=body_font, fill="#25221e")
        y += 105
    draw.line((220, y + 30, width_px - 220, y + 30), fill="#9a9183", width=2)
    draw.text((220, y + 90), "记录人：演示人员    日期：2026-06-18", font=small_font, fill="#514c45")
    draw.text((220, height_px - 240), "演示扫描件 - 不作为真实地质成果", font=small_font, fill="#7b332c")
    image = image.filter(ImageFilter.GaussianBlur(radius=0.35))
    scan_png = TMP / "f1-field-note-scan.png"
    image.save(scan_png, format="PNG", dpi=(300, 300))
    output = DEMO / "04-F1断裂野外记录扫描件.pdf"
    pdf = canvas.Canvas(str(output), pagesize=A4)
    pdf.setTitle("F1断裂野外记录扫描件（演示数据）")
    pdf.drawImage(str(scan_png), 0, 0, width=A4[0], height=A4[1])
    pdf.showPage()
    pdf.save()
    OUTPUT_PDF.mkdir(parents=True, exist_ok=True)
    (OUTPUT_PDF / output.name).write_bytes(output.read_bytes())


def main() -> None:
    DEMO.mkdir(parents=True, exist_ok=True)
    DOCS.mkdir(parents=True, exist_ok=True)
    build_drill_docx()
    build_searchable_pdf()
    build_scanned_pdf()
    markdown_to_docx(
        DOCS / "demo-tutorial.md",
        DOCS / "系统演示操作指南.docx",
        "系统演示操作指南",
        "基于填图对象智能识别的文本空间化算法模块",
        "v1.1 优化演示数据版 | 15-20分钟 | 项目验收与培训",
        "editorial_cover",
    )
    markdown_to_docx(
        DOCS / "project-demo-speech.md",
        DOCS / "项目功能演示讲稿.docx",
        "项目功能演示讲稿",
        "逐步操作、逐字话术与答辩备用问答",
        "v1.1 | 正式演示版",
        "workshop_agenda",
    )
    print("Demo package generated:")
    for path in sorted(DEMO.iterdir()):
        print(path.name)
    print((DOCS / "系统演示操作指南.docx").name)
    print((DOCS / "项目功能演示讲稿.docx").name)


if __name__ == "__main__":
    main()

from io import BytesIO

from docx import Document

from app.parsers.base import RawDocument, RawParagraph


class WordParser:
    def parse(self, content: bytes) -> RawDocument:
        document = Document(BytesIO(content))
        paragraphs: list[RawParagraph] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
            paragraphs.append(RawParagraph(1, text, style_name.startswith("heading") or style_name.startswith("标题")))
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(RawParagraph(1, " | ".join(cells)))
        return RawDocument("WORD", 1, paragraphs)


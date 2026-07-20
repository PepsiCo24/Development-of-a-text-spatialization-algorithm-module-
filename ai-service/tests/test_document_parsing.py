from io import BytesIO

import fitz
from docx import Document

from app.parsers.pipeline import DocumentParsingPipeline


class StubOcr:
    def recognize(self, image_bytes: bytes) -> str:
        return "第一章 区域地质\n燕山期花岗岩出露于大冶地区。"


def test_plain_text_is_cleaned_structured_and_chunked() -> None:
    source = "第一章 区域地质\n\n  花岗岩   主要分布于大冶。\n\n1.1 构造\n\n北东向断裂控制矿体。".encode()
    result = DocumentParsingPipeline(ocr=StubOcr(), chunk_size=40).parse("report.txt", source)

    assert result.document_type == "TXT"
    assert result.page_count == 1
    assert len(result.chunks) == 2
    assert result.chunks[0].chapter_title == "第一章 区域地质"
    assert result.chunks[0].content == "花岗岩 主要分布于大冶。"
    assert result.chunks[1].chapter_title == "1.1 构造"
    assert "断裂控制矿体" in result.chunks[1].content


def test_docx_headings_and_tables_are_extracted() -> None:
    document = Document()
    document.add_heading("第一章 矿区概况", level=1)
    document.add_paragraph("矿区位于湖北省黄石市。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "矿种"
    table.cell(0, 1).text = "铁矿"
    output = BytesIO()
    document.save(output)

    result = DocumentParsingPipeline(ocr=StubOcr()).parse("report.docx", output.getvalue())

    assert result.document_type == "WORD"
    assert result.chunks[0].chapter_title == "第一章 矿区概况"
    assert "矿区位于湖北省黄石市" in result.chunks[0].content
    assert "矿种 | 铁矿" in result.chunks[0].content


def test_pdf_text_blocks_preserve_page_numbers() -> None:
    document = fitz.open()
    first = document.new_page()
    first.insert_text((72, 72), "Chapter One")
    first.insert_text((72, 100), "Granite body in Daye district with iron ore.")
    second = document.new_page()
    second.insert_text((72, 72), "Second page describes a northeast fault zone.")
    content = document.tobytes()
    document.close()

    result = DocumentParsingPipeline(ocr=StubOcr()).parse("report.pdf", content)

    assert result.document_type == "PDF"
    assert result.page_count == 2
    assert result.chunks[0].page_start == 1
    assert result.chunks[-1].page_end == 2
    assert "Granite body" in " ".join(chunk.content for chunk in result.chunks)


def test_image_uses_ocr_service() -> None:
    result = DocumentParsingPipeline(ocr=StubOcr()).parse("scan.png", b"image-bytes-are-owned-by-stub")

    assert result.document_type == "IMAGE"
    assert result.warnings == ["图片使用 PaddleOCR 识别"]
    assert result.chunks[0].chapter_title == "第一章 区域地质"
    assert "花岗岩" in result.chunks[0].content

